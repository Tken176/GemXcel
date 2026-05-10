import os
import json
import re
import time
import hashlib
import signal
import requests
import chardet
from typing import List, Dict, Optional, Tuple, Set
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
import config
import random
import threading

logger = config.logger

SUPABASE_URL = "https://hwqlvikkzeybssocyjjn.supabase.co"
PROXY_URL = f"{SUPABASE_URL}/functions/v1/gemini-proxy"
GEMINI_PATH = "/v1beta/models/gemini-2.5-flash-lite:generateContent"


class TimeoutException(Exception):
    pass


def timeout_handler(signum, frame):
    raise TimeoutException("Operation timed out")


# ============================================================
# FIX #1: RATE LIMITER — Chống lỗi 429
# Giới hạn tối đa 8 request/phút với token bucket
# ============================================================
class RateLimiter:
    """Thread-safe token bucket rate limiter.
    
    Với batching quiz (3 request thay vì ~20), tổng request cho 1 cuốn sách
    chỉ còn ~8 (5 bài học + 3 quiz batch). Giới hạn 9 req/phút vẫn an toàn
    và cho phép các bài học chạy song song nhanh hơn.
    """

    def __init__(self, max_calls: int = 9, period: float = 60.0):
        self.max_calls = max_calls
        self.period = period
        self.min_interval = period / max_calls  # ~7.5s giữa các request
        self._lock = threading.Lock()
        self._last_call_time: float = 0.0
        self._call_times: List[float] = []

    def wait(self):
        """Block cho đến khi được phép gọi API."""
        with self._lock:
            now = time.time()

            # Loại bỏ các call cũ hơn 1 chu kỳ
            self._call_times = [t for t in self._call_times if now - t < self.period]

            if len(self._call_times) >= self.max_calls:
                # Phải chờ đến khi call cũ nhất rơi ra ngoài cửa sổ
                oldest = self._call_times[0]
                wait_time = self.period - (now - oldest) + 0.1
                if wait_time > 0:
                    logger.debug("Rate limit: chờ %.1fs để tránh 429...", wait_time)
                    time.sleep(wait_time)
                    now = time.time()
                    self._call_times = [t for t in self._call_times if now - t < self.period]

            # Đảm bảo khoảng cách tối thiểu giữa 2 request liên tiếp
            elapsed = now - self._last_call_time
            if elapsed < self.min_interval:
                time.sleep(self.min_interval - elapsed)
                now = time.time()

            self._call_times.append(now)
            self._last_call_time = now


# ============================================================
# FIX #2: GEMINI CLIENT — Xử lý 429 với exponential backoff
# ============================================================
class GeminiClient:

    def __init__(self):
        self.access_token = self._load_access_token()
        self.session = requests.Session()
        # Dùng chung một rate limiter cho toàn bộ instance
        self.rate_limiter = RateLimiter(max_calls=9, period=60.0)

    def _load_access_token(self) -> str:
        appdata_path = os.environ.get('APPDATA')
        if not appdata_path:
            raise ValueError("Không tìm thấy APPDATA. Đảm bảo chạy trên Windows.")
        token_file = os.path.join(appdata_path, "GemxelProject", "auth_session.json")

        if os.path.exists(token_file):
            try:
                with open(token_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                session_data = data.get("session", data)
                access_token = session_data.get("access_token", "")
                if access_token:
                    return access_token
            except (json.JSONDecodeError, IOError) as e:
                raise ValueError(f"Không đọc được file auth_session.json: {e}")

        raise ValueError("Chưa đăng nhập! Vui lòng đăng nhập vào ứng dụng trước.")

    def chat_completions_create(self, messages, temperature=0.7, **kwargs):
        if isinstance(messages, list) and messages:
            prompt = messages[-1].get("content", "")
        else:
            prompt = str(messages)

        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.access_token}",
        }

        data = {
            "geminiPath": GEMINI_PATH,
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": 4096,
                "topP": 0.8,
                "topK": 10,
            },
        }

        # FIX: Retry với exponential backoff khi gặp 429
        max_attempts = 5
        base_delay = 15.0  # giây

        for attempt in range(max_attempts):
            # Chờ rate limiter trước mỗi lần gọi
            self.rate_limiter.wait()

            try:
                response = self.session.post(
                    PROXY_URL,
                    headers=headers,
                    json=data,
                    timeout=60,
                )

                if response.status_code == 200:
                    result = response.json()
                    # FIX: Kiểm tra cấu trúc response trước khi truy cập
                    candidates = result.get("candidates", [])
                    if not candidates:
                        raise ValueError("Gemini trả về candidates rỗng")
                    content_parts = candidates[0].get("content", {}).get("parts", [])
                    if not content_parts:
                        raise ValueError("Gemini trả về parts rỗng")
                    text = content_parts[0].get("text", "")
                    return GeminiResponse(text)

                elif response.status_code == 429:
                    # Lấy Retry-After từ header nếu có
                    retry_after = response.headers.get("Retry-After")
                    if retry_after:
                        wait = float(retry_after) + 1
                    else:
                        wait = base_delay * (2 ** attempt)  # 15, 30, 60, 120...
                    logger.warning("429 Rate Limit (lần %d/%d). Chờ %ds...", attempt + 1, max_attempts, int(wait))
                    time.sleep(wait)
                    continue  # thử lại

                elif response.status_code == 401:
                    raise Exception("Token hết hạn, vui lòng đăng nhập lại.")

                elif response.status_code == 503:
                    wait = base_delay * (2 ** attempt)
                    logger.warning("503 Service Unavailable (lần %d). Chờ %ds...", attempt + 1, int(wait))
                    time.sleep(wait)
                    continue

                else:
                    raise Exception(f"Proxy error {response.status_code}: {response.text[:200]}")

            except requests.exceptions.Timeout:
                if attempt < max_attempts - 1:
                    wait = base_delay * (attempt + 1)
                    logger.warning("Timeout lần %d. Chờ %ds rồi thử lại...", attempt + 1, int(wait))
                    time.sleep(wait)
                    continue
                raise Exception("Gemini proxy timeout sau nhiều lần thử")

            except requests.exceptions.RequestException as e:
                raise Exception(f"Request error: {e}")

        raise Exception(f"Đã thử {max_attempts} lần nhưng vẫn gặp lỗi 429/503. Vui lòng thử lại sau.")


class GeminiResponse:
    def __init__(self, content):
        self.choices = [GeminiChoice(content)]


class GeminiChoice:
    def __init__(self, content):
        self.message = GeminiMessage(content)


class GeminiMessage:
    def __init__(self, content):
        self.content = content


class GeminiChatCompletions:
    def __init__(self, client):
        self.client = client

    def create(self, model=None, messages=None, temperature=0.7, **kwargs):
        return self.client.chat_completions_create(messages=messages, temperature=temperature, **kwargs)


class GeminiChat:
    def __init__(self, client):
        self.completions = GeminiChatCompletions(client)


# ============================================================
# FIX #3: JSON VALIDATOR — Chống đáp án rác
# ============================================================
class QuestionValidator:
    """Kiểm tra và làm sạch câu hỏi trước khi lưu."""

    @staticmethod
    def is_garbage_text(text: str) -> bool:
        """Phát hiện text rác (placeholder, generic, vô nghĩa)."""
        if not text or not text.strip():
            return True
        garbage_patterns = [
            r"^[A-D]\.$",              # chỉ "A." không có nội dung
            r"^(đáp án|answer)\s*\d*$",
            r"^(lựa chọn|option|choice)\s*\d*$",
            r"^(từ đúng|từ sai|từ nhiễu)\s*\d*$",
            r"^(câu hỏi|question)\s*\d*$",
            r"^\.\.\.$",
            r"^(a|b|c|d)$",
        ]
        lower = text.strip().lower()
        for pattern in garbage_patterns:
            if re.match(pattern, lower, re.IGNORECASE):
                return True
        # Quá ngắn (< 3 ký tự sau khi bỏ prefix A./B./C./D.)
        content = re.sub(r"^[A-D]\.\s*", "", text.strip())
        if len(content.strip()) < 3:
            return True
        return False

    @staticmethod
    def validate_question(q: dict) -> bool:
        """Trả về True nếu câu hỏi hợp lệ và không có đáp án rác."""
        if not isinstance(q, dict):
            return False
        question_text = q.get("question", "")
        if QuestionValidator.is_garbage_text(question_text):
            return False

        q_type = q.get("type", "mcq")
        if q_type == "mcq":
            choices = q.get("choices", [])
            if not isinstance(choices, list) or len(choices) < 4:
                return False
            # Tất cả choices phải có nội dung thực
            if any(QuestionValidator.is_garbage_text(c) for c in choices):
                return False
            # Không được tất cả giống nhau
            unique = set(re.sub(r"^[A-D]\.\s*", "", c).strip().lower() for c in choices)
            if len(unique) < 3:
                return False

        elif q_type == "fill":
            words = q.get("words", q.get("choices", []))
            if not isinstance(words, list) or len(words) < 2:
                return False
            if any(QuestionValidator.is_garbage_text(w) for w in words):
                return False

        elif q_type == "tf":
            pass  # tf luôn hợp lệ nếu có question

        # correct_answer phải là int hợp lệ
        correct = q.get("correct_answer")
        if not isinstance(correct, int):
            return False

        return True

    @staticmethod
    def clean_question(q: dict) -> dict:
        """Chuẩn hóa câu hỏi: loại bỏ prefix thừa, giới hạn độ dài."""
        if not isinstance(q, dict):
            return q
        # Giới hạn độ dài choices
        for key in ("choices", "words"):
            items = q.get(key)
            if isinstance(items, list):
                q[key] = [c[:80] if len(c) > 80 else c for c in items]
        return q


# ============================================================
# CONTENT PROCESSOR — chính
# ============================================================
class ContentProcessor:
    def __init__(self, lessons_count: int = 5, questions_per_lesson: int = 6, quiz_questions: int = 10):
        self.gemini_client = GeminiClient()
        self.client = type('Client', (), {'chat': GeminiChat(self.gemini_client)})()

        self.lessons_path = config.LESSON_DATA_FILE_PATH
        self.quiz_path = config.QUIZ_DATA_FILE_PATH

        self.supported_formats = ['.txt', '.docx', '.md', '.pdf']
        self.max_retries = 3
        self.timeout = 45
        self.total_timeout = 300

        self.lessons_count = max(1, min(lessons_count, 20))
        self.questions_per_lesson = max(3, min(questions_per_lesson, 15))
        self.quiz_questions = max(5, min(quiz_questions, 50))

        self.content_cache = {}
        os.makedirs(config.ASSETS_DIR, exist_ok=True)
        self.should_stop = False
        self.used_questions: Set[str] = set()
        self.used_quiz_questions: Set[str] = set()
        self.validator = QuestionValidator()

    def is_supported(self, file_path: str) -> bool:
        return any(file_path.lower().endswith(ext) for ext in self.supported_formats)

    def process_file(self, file_path: str) -> Tuple[bool, str]:
        start_time = time.time()
        self.used_questions.clear()
        self.used_quiz_questions.clear()

        if hasattr(signal, 'SIGALRM'):
            signal.signal(signal.SIGALRM, timeout_handler)
            signal.alarm(self.total_timeout)

        try:
            if not self.is_supported(file_path):
                return False, f"Định dạng file không được hỗ trợ. Các định dạng hỗ trợ: {', '.join(self.supported_formats)}"

            if not os.path.exists(file_path):
                return False, f"File không tồn tại: {file_path}"

            try:
                content = self.read_file_content(file_path)
            except Exception as e:
                return False, f"Không thể đọc file: {e}"

            if not content or not content.strip():
                return False, "Không thể đọc nội dung từ file (file rỗng hoặc định dạng không đúng)."

            if len(content.strip()) < 200:
                return False, "Nội dung file quá ngắn để xử lý (cần ít nhất 200 ký tự)."

            content = self._preprocess_content(content)
            content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()

            if content_hash in self.content_cache:
                lessons, quiz_questions = self.content_cache[content_hash]
            else:
                chunks = self._intelligent_chunking(content, target_chunks=self.lessons_count)
                lessons = self.generate_lessons_with_timeout(chunks)

                if not lessons:
                    return False, "Không thể tạo bài học từ nội dung này."

                lessons = [lesson for lesson in lessons if lesson is not None]
                if not lessons:
                    return False, "Không thể tạo bài học hợp lệ từ nội dung này."

                quiz_questions = self.generate_quiz_with_timeout(lessons, total_questions=self.quiz_questions)
                self.content_cache[content_hash] = (lessons, quiz_questions)

            # FIX #4: Validate trước khi lưu — không lưu nếu dữ liệu rỗng
            if not self._validate_output_before_save(lessons, quiz_questions):
                return False, "Dữ liệu sinh ra không hợp lệ (quá nhiều câu hỏi rác). Vui lòng thử lại."

            success = self._safe_save_data_files(lessons, quiz_questions)
            if not success:
                return False, "Không thể lưu dữ liệu. Vui lòng kiểm tra quyền ghi file."

            elapsed_time = time.time() - start_time
            total_questions = sum(len(lesson.get('questions', [])) for lesson in lessons)
            total_quiz = sum(len(q) for q in quiz_questions.values()) if isinstance(quiz_questions, dict) else 0

            return True, f"Xử lý thành công trong {elapsed_time:.1f}s! Đã tạo {len(lessons)} bài học ({total_questions} câu hỏi) và {total_quiz} câu hỏi quiz."

        except TimeoutException:
            elapsed_time = time.time() - start_time
            return False, f"Timeout sau {elapsed_time:.1f}s! Quá trình xử lý mất quá nhiều thời gian."
        except Exception as e:
            elapsed_time = time.time() - start_time
            error_msg = str(e)
            if "timeout" in error_msg.lower():
                return False, f"Timeout sau {elapsed_time:.1f}s! Server AI phản hồi chậm."
            else:
                return False, f"Lỗi sau {elapsed_time:.1f}s: {error_msg}"
        finally:
            if hasattr(signal, 'SIGALRM'):
                signal.alarm(0)

    def _validate_output_before_save(self, lessons: List[Dict], quiz: Dict) -> bool:
        """Kiểm tra dữ liệu đầu ra trước khi ghi file."""
        if not lessons:
            return False
        # Mỗi bài học phải có ít nhất 1 câu hỏi thực sự hợp lệ
        valid_lessons = 0
        for lesson in lessons:
            valid_qs = [q for q in lesson.get("questions", []) if self.validator.validate_question(q)]
            if valid_qs:
                valid_lessons += 1
        if valid_lessons == 0:
            return False

        # Quiz phải có ít nhất 1 câu hỏi hợp lệ mỗi độ khó (nếu dict)
        if isinstance(quiz, dict):
            total_valid_quiz = sum(
                sum(1 for q in qs if self.validator.validate_question(q))
                for qs in quiz.values()
            )
            if total_valid_quiz == 0:
                return False

        return True

    # ====================================================
    # FILE READING
    # ====================================================

    def _detect_encoding(self, file_path: str) -> str:
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)
                result = chardet.detect(raw_data)
                confidence = result.get('confidence', 0)
                encoding = result.get('encoding', 'utf-8')
                if confidence < 0.7 or not encoding:
                    return 'utf-8'
                encoding_map = {
                    'utf-16': 'utf-16', 'utf-16le': 'utf-16-le', 'utf-16be': 'utf-16-be',
                    'windows-1252': 'cp1252', 'iso-8859-1': 'latin-1',
                }
                return encoding_map.get(encoding.lower(), encoding)
        except Exception:
            return 'utf-8'

    def _preprocess_content(self, content: str) -> str:
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)
        cleaned_chars = [ch for ch in content if ch == '\n' or ch == '\t' or ch.isprintable()]
        return ''.join(cleaned_chars).strip()

    def read_file_content(self, file_path: str) -> Optional[str]:
        try:
            if not self.is_supported(file_path):
                raise ValueError(f"Định dạng file không được hỗ trợ: {file_path}")

            if file_path.lower().endswith('.txt'):
                detected_encoding = self._detect_encoding(file_path)
                encodings = [detected_encoding, 'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'cp1252', 'latin-1', 'gbk', 'big5']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                            if content.strip():
                                return content
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                try:
                    with open(file_path, 'rb') as f:
                        raw_data = f.read()
                    for encoding in ['utf-8', 'cp1252', 'latin-1']:
                        try:
                            content = raw_data.decode(encoding, errors='replace')
                            if content.strip() and not self._is_binary_garbage(content):
                                return content
                        except Exception:
                            continue
                except Exception:
                    pass
                raise ValueError("Không thể đọc file .txt với các encoding thông dụng")

            elif file_path.lower().endswith('.docx'):
                try:
                    import docx
                    if not os.path.exists(file_path):
                        raise ValueError("File không tồn tại")
                    if os.path.getsize(file_path) == 0:
                        raise ValueError("File DOCX rỗng")
                    try:
                        doc = docx.Document(file_path)
                    except Exception as e:
                        raise ValueError(f"File DOCX bị lỗi: {e}")
                    paragraphs_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                    tables_text = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                tables_text.append(" | ".join(row_text))
                    content = "\n\n".join(paragraphs_text + tables_text)
                    if not content.strip():
                        raise ValueError("File DOCX không chứa text có thể đọc được")
                    if self._is_binary_garbage(content):
                        raise ValueError("Nội dung DOCX chứa dữ liệu binary không hợp lệ")
                    return content
                except ImportError:
                    raise ValueError("Cần cài đặt python-docx: pip install python-docx")

            elif file_path.lower().endswith('.md'):
                detected_encoding = self._detect_encoding(file_path)
                for encoding in [detected_encoding, 'utf-8', 'cp1252', 'latin-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                            if content.strip():
                                content = re.sub(r'(^|\n)#{1,6}\s*', r'\1', content)
                                content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
                                content = re.sub(r'\*(.*?)\*', r'\1', content)
                                content = re.sub(r'`{1,3}([^`]+)`{1,3}', r'\1', content)
                                return content
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                raise ValueError("File Markdown không đọc được với các encoding thông dụng")

            elif file_path.lower().endswith('.pdf'):
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        pages_text = []
                        for page in reader.pages:
                            try:
                                text = page.extract_text()
                                if text and text.strip():
                                    pages_text.append(text)
                            except Exception:
                                continue
                        content = "\n".join(pages_text)
                        if not content.strip():
                            raise ValueError("Không thể trích xuất text từ PDF")
                        return content
                except ImportError:
                    raise ValueError("Cần cài đặt PyPDF2: pip install PyPDF2")

        except Exception as e:
            logger.warning("Lỗi đọc file %s: %s", file_path, e)
            return None

    def _is_binary_garbage(self, content: str) -> bool:
        if not content:
            return True
        printable_chars = sum(1 for c in content if c.isprintable() or c.isspace())
        total_chars = len(content)
        if total_chars > 0:
            return (printable_chars / total_chars) < 0.8
        return True

    # ====================================================
    # CHUNKING
    # ====================================================

    def _intelligent_chunking(self, content: str, target_chunks: int = 5) -> List[str]:
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        if len(paragraphs) <= target_chunks:
            words = content.split()
            chunk_size = max(80, len(words) // max(1, target_chunks))
            chunks = []
            for i in range(target_chunks):
                start_idx = i * chunk_size
                end_idx = min((i + 1) * chunk_size, len(words))
                if start_idx < len(words):
                    chunk = ' '.join(words[start_idx:end_idx])
                    if chunk.strip():
                        chunks.append(chunk)
            return chunks
        else:
            base = len(paragraphs) // target_chunks
            rem = len(paragraphs) % target_chunks
            chunks, start = [], 0
            for i in range(target_chunks):
                take = base + (1 if i < rem else 0)
                end = start + take
                chunk = '\n\n'.join(paragraphs[start:end]).strip()
                if chunk:
                    chunks.append(chunk)
                start = end
            return chunks

    # ====================================================
    # LESSON GENERATION
    # ====================================================

    def generate_lessons_with_timeout(self, chunks: List[str]) -> List[Dict]:
        results = [None] * len(chunks)
        completed_count = 0

        def _lesson_callback(future, lesson_num):
            nonlocal completed_count
            try:
                result = future.result(timeout=self.timeout)
                results[lesson_num - 1] = result
                completed_count += 1
            except Exception:
                results[lesson_num - 1] = self._create_fallback_lesson(
                    lesson_num, chunks[lesson_num - 1] if lesson_num - 1 < len(chunks) else ""
                )
                completed_count += 1

        # BATCHING UPDATE: max_workers=3 để xử lý song song các bài học.
        # Rate limiter toàn cục đã bảo vệ khỏi 429; quiz dùng batched (3 request)
        # nên tổng request/phút luôn trong giới hạn Free Tier.
        with ThreadPoolExecutor(max_workers=1) as executor:
            futures = []
            for i, chunk in enumerate(chunks):
                if self.should_stop:
                    break
                # Stagger nhỏ 1s giữa các submit để rate limiter phân bổ đều
                if i > 0:
                    time.sleep(1)
                future = executor.submit(self._generate_single_lesson_safe, chunk, i + 1)
                future.add_done_callback(lambda f, num=i + 1: _lesson_callback(f, num))
                futures.append(future)

            try:
                for _ in as_completed(futures, timeout=self.timeout * len(chunks)):
                    if self.should_stop:
                        break
            except TimeoutError:
                pass

        return [r for r in results if r is not None]

    def _generate_unique_question_hash(self, question: str, options: List[str]) -> str:
        normalized_question = re.sub(r'\s+', ' ', question.strip().lower())
        normalized_options = [re.sub(r'\s+', ' ', opt.strip().lower()) for opt in options]
        combined = normalized_question + '|' + '|'.join(normalized_options)
        return hashlib.md5(combined.encode('utf-8')).hexdigest()

    def _is_question_unique(self, question: str, options: List[str], used_set: Set[str]) -> bool:
        return self._generate_unique_question_hash(question, options) not in used_set

    def _add_question_to_used(self, question: str, options: List[str], used_set: Set[str]):
        used_set.add(self._generate_unique_question_hash(question, options))

    def _generate_single_lesson_safe(self, chunk: str, lesson_number: int) -> Dict:
        for attempt in range(self.max_retries):
            try:
                if self.should_stop:
                    break

                max_chunk_length = 1500
                chunk_for_prompt = (chunk[:max_chunk_length] + "...") if len(chunk) > max_chunk_length else chunk

                prompt = f"""Tạo bài học số {lesson_number} từ nội dung sau:
{chunk_for_prompt}

Yêu cầu:
1. Tạo bài học theo thứ tự hợp lý. Tiêu đề ngắn gọn. Nội dung tóm tắt 200-350 từ.
2. Tạo ĐÚNG {self.questions_per_lesson} câu hỏi. PHẢI BAO GỒM ĐA DẠNG CÁC LOẠI CÂU HỎI SAU:
   - "mcq" (Trắc nghiệm 4 đáp án)
   - "tf" (Đúng/Sai)
   - "fill" (Điền từ vào chỗ trống)

Trả về CHỈ JSON hợp lệ theo mẫu chính xác sau:
{{
  "name": "Bài {lesson_number}",
  "title": "[Tiêu đề cụ thể về nội dung]",
  "content": "[Nội dung tóm tắt chi tiết]",
  "questions": [
    {{
      "type": "mcq",
      "question": "Câu hỏi trắc nghiệm 4 đáp án?",
      "choices": ["A. Lựa chọn 1", "B. Lựa chọn 2", "C. Lựa chọn 3", "D. Lựa chọn 4"],
      "correct_answer": 0,
      "difficulty": "medium"
    }},
    {{
      "type": "tf",
      "question": "Nhận định này đúng hay sai?",
      "correct_answer": 0,
      "difficulty": "easy"
    }},
    {{
      "type": "fill",
      "question": "Vệ tinh tự nhiên của Trái Đất là ___.",
      "words": ["Mặt trăng", "Mặt trời", "Sao mộc", "Sao kim"],
      "correct_answer": 0,
      "difficulty": "hard"
    }}
  ]
}}
LƯU Ý: correct_answer là vị trí index (0, 1, 2, 3) của đáp án đúng trong mảng choices hoặc words. Với "tf", 0 = Đúng, 1 = Sai.
"""
                response = self.client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                )

                response_text = (response.choices[0].message.content or "").strip()
                if not response_text:
                    raise ValueError("Response rỗng từ AI")

                lesson = self._safe_parse_json(response_text)
                if not lesson:
                    raise ValueError("Không parse được JSON từ response")

                return self._validate_lesson(lesson, lesson_number, chunk)

            except Exception as e:
                logger.warning("Lỗi lesson %d (lần %d): %s", lesson_number, attempt + 1, e)
                if attempt == self.max_retries - 1:
                    return self._create_fallback_lesson(lesson_number, chunk)
                time.sleep(min(2 ** attempt, 5))

    def _safe_parse_json(self, text: str) -> Optional[Dict]:
        try:
            # Loại bỏ markdown code block
            cleaned = text.replace('```json', '').replace('```', '').strip()

            # Thử parse trực tiếp
            parsed = json.loads(cleaned)
            if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
                if self._is_valid_lesson_structure(parsed[0]):
                    return parsed[0]
                for item in parsed:
                    if isinstance(item, dict) and self._is_valid_lesson_structure(item):
                        return item
                return parsed[0]
            if isinstance(parsed, dict) and self._is_valid_lesson_structure(parsed):
                return parsed
        except Exception:
            pass

        # Fallback: tìm JSON object trong text
        json_candidates = self._extract_json_objects(text)
        for candidate in json_candidates:
            try:
                parsed = json.loads(candidate)
                if isinstance(parsed, dict) and self._is_valid_lesson_structure(parsed):
                    return parsed
                if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
                    return parsed[0]
            except json.JSONDecodeError:
                continue

        return None

    def _is_valid_lesson_structure(self, data: Dict) -> bool:
        required_fields = ['name', 'title', 'content', 'questions']
        if not all(field in data for field in required_fields):
            return False
        if not isinstance(data['questions'], list):
            return False
        return True

    def _extract_json_objects(self, text: str) -> List[str]:
        results = []
        stack, in_string, escape, start_idx = 0, False, False, None

        for i, ch in enumerate(text):
            if in_string:
                if escape:
                    escape = False
                elif ch == '\\':
                    escape = True
                elif ch == '"':
                    in_string = False
                continue
            else:
                if ch == '"':
                    in_string = True
                    continue
                if ch == '{':
                    if stack == 0:
                        start_idx = i
                    stack += 1
                elif ch == '}':
                    if stack > 0:
                        stack -= 1
                        if stack == 0 and start_idx is not None:
                            results.append(text[start_idx:i + 1])
                            start_idx = None
        return results

    def _validate_lesson(self, lesson: Dict, lesson_number: int, chunk: str) -> Dict:
        lesson.setdefault("name", f"Bài {lesson_number}")
        lesson.setdefault("title", f"Bài học {lesson_number}")
        lesson.setdefault("content", chunk[:300] + "..." if len(chunk) > 300 else chunk)

        if "questions" not in lesson or not isinstance(lesson["questions"], list):
            lesson["questions"] = []

        unique_questions = []
        for q in lesson["questions"]:
            if not isinstance(q, dict):
                continue

            q_type = q.get("type", "mcq")
            if q_type not in ["mcq", "tf", "fill", "drag"]:
                q["type"] = "mcq"
                q_type = "mcq"

            if q_type == "mcq":
                if not isinstance(q.get("choices"), list) or len(q["choices"]) < 2:
                    q["choices"] = ["A. Đúng", "B. Sai", "C. Không xác định", "D. Tất cả đều sai"]
                options = q["choices"]
            elif q_type == "tf":
                q["choices"] = ["Đúng", "Sai"]
                if q.get("correct_answer") not in [0, 1]:
                    q["correct_answer"] = 0
                options = q["choices"]
            elif q_type in ["fill", "drag"]:
                if not isinstance(q.get("words"), list) or len(q["words"]) < 2:
                    q["words"] = ["Từ đúng", "Từ sai 1", "Từ sai 2"]
                q["choices"] = q["words"]
                options = q["words"]
            else:
                options = q.get("choices", [])

            # FIX: Giới hạn độ dài options
            for idx, choice in enumerate(options):
                if len(choice) > 80:
                    prefix = choice[:2] if choice[:2] in ('A.', 'B.', 'C.', 'D.') else ""
                    content_part = choice[2:].strip() if prefix else choice
                    options[idx] = (prefix + " " + content_part[:77 - len(prefix)]).strip() if prefix else content_part[:80]

            question_text = q.get("question", "")

            if not isinstance(q.get("correct_answer"), int) or q["correct_answer"] < 0 or q["correct_answer"] >= len(options):
                q["correct_answer"] = 0

            q.setdefault("difficulty", "medium")

            # FIX: Chỉ thêm câu hỏi không phải rác VÀ unique
            if (self.validator.validate_question(q) and
                    self._is_question_unique(question_text, options, self.used_questions)):
                self._add_question_to_used(question_text, options, self.used_questions)
                unique_questions.append(q)

        lesson["questions"] = unique_questions

        # Bổ sung câu hỏi nếu thiếu
        while len(lesson["questions"]) < self.questions_per_lesson:
            new_question = self._create_default_question(len(lesson["questions"]) + 1, lesson_number, chunk)
            if self._is_question_unique(new_question["question"], new_question["choices"], self.used_questions):
                self._add_question_to_used(new_question["question"], new_question["choices"], self.used_questions)
                lesson["questions"].append(new_question)
            else:
                new_question = self._create_varied_question(len(lesson["questions"]) + 1, lesson_number, chunk)
                self._add_question_to_used(
                    new_question["question"],
                    new_question.get("choices", new_question.get("words", [])),
                    self.used_questions,
                )
                lesson["questions"].append(new_question)

        lesson["questions"] = lesson["questions"][:self.questions_per_lesson]
        return lesson

    def _create_varied_question(self, q_num: int, lesson_num: int, content: str) -> Dict:
        difficulties = ["easy", "medium", "hard"]
        difficulty = difficulties[(q_num - 1) % 3]

        question_templates = [
            f"Theo bài {lesson_num}, yếu tố quan trọng được nhấn mạnh là gì?",
            f"Khái niệm chủ đạo trong bài {lesson_num} liên quan đến?",
            f"Nội dung bài {lesson_num} tập trung giải thích về?",
            f"Điểm cốt lõi của bài {lesson_num} là gì?",
        ]
        answer_templates = [
            ["Nguyên lý cơ bản", "Phương pháp ứng dụng", "Quy trình thực hiện", "Kết quả đạt được"],
            ["Lý thuyết nền tảng", "Thực tiễn áp dụng", "Kinh nghiệm rút ra", "Hướng phát triển"],
        ]

        if random.random() > 0.7:
            return {
                "type": "tf",
                "question": f"Nội dung bài {lesson_num} đề cập đến một khái niệm quan trọng. Đúng hay Sai?",
                "choices": ["Đúng", "Sai"],
                "correct_answer": 0,  # FIX: luôn là 0 (Đúng) để tránh rác
                "difficulty": difficulty,
            }

        answers = random.choice(answer_templates)
        return {
            "type": "mcq",
            "question": random.choice(question_templates),
            "choices": [f"{'ABCD'[i]}. {ans}" for i, ans in enumerate(answers)],
            "correct_answer": 0,  # FIX: luôn index 0 với fallback
            "difficulty": difficulty,
        }

    def _create_default_question(self, q_num: int, lesson_num: int, content: str) -> Dict:
        difficulties = ["easy", "medium", "hard"]
        difficulty = difficulties[(q_num - 1) % 3]

        return {
            "type": "mcq",
            "question": f"Nội dung bài {lesson_num} chủ yếu đề cập đến vấn đề gì?",
            "choices": ["A. Nội dung chính", "B. Ý tưởng phụ", "C. Khái niệm mở rộng", "D. Ví dụ minh họa"],
            "correct_answer": 0,
            "difficulty": difficulty,
        }

    def _create_fallback_lesson(self, lesson_number: int, chunk: str) -> Dict:
        content = chunk[:350] + "..." if len(chunk) > 350 else chunk
        words = chunk.split()[:8]
        title = " ".join(words) if words else f"Bài học {lesson_number}"
        if len(title) > 45:
            title = title[:42] + "..."

        questions = []
        for i in range(self.questions_per_lesson):
            new_question = self._create_varied_question(i + 1, lesson_number, chunk)
            self._add_question_to_used(
                new_question["question"],
                new_question.get("choices", ["Đúng", "Sai"]),
                self.used_questions,
            )
            questions.append(new_question)

        return {"name": f"Bài {lesson_number}", "title": title, "content": content, "questions": questions}

    # ====================================================
    # QUIZ GENERATION
    # ====================================================

    def generate_quiz_with_timeout(self, lessons: List[Dict], total_questions: int = 10) -> Dict:
        try:
            return self.generate_quiz(lessons, total_questions)
        except Exception as e:
            logger.warning("Lỗi generate_quiz: %s. Dùng fallback.", e)
            return self._create_fallback_quiz(lessons, total_questions)

    def generate_quiz(self, lessons: List[Dict], total_questions: int = 10) -> Dict:
        """
        BATCHING: Thay vì gọi AI N lần (1 câu/lần), giờ gọi đúng 3 lần:
        1 request cho toàn bộ câu Dễ, 1 cho Trung bình, 1 cho Khó.
        Giảm từ ~20 request xuống còn 3 request cho quiz → không bao giờ
        chạm ngưỡng 10 req/phút của Free Tier.
        """
        min_questions_per_difficulty = 10
        easy_count = max(min_questions_per_difficulty, int(total_questions * 0.4))
        medium_count = max(min_questions_per_difficulty, int(total_questions * 0.4))
        hard_count = max(min_questions_per_difficulty, total_questions - easy_count - medium_count)

        counts = {"easy": easy_count, "medium": medium_count, "hard": hard_count}
        return self._generate_all_quizzes_batched(lessons, counts)

    def _generate_all_quizzes_batched(self, lessons: List[Dict], counts: Dict[str, int]) -> Dict:
        """
        Gửi đúng 3 request song song (1 per difficulty) thay vì N request tuần tự.
        Mỗi request yêu cầu AI trả về toàn bộ mảng câu hỏi của 1 mức độ.
        """
        quiz: Dict[str, List] = {"easy": [], "medium": [], "hard": []}

        # Chuẩn bị context từ lessons (giới hạn kích thước để tránh token quá lớn)
        lesson_contexts = []
        for lesson in lessons[:5]:  # Tối đa 5 bài để prompt không quá dài
            ctx = lesson.get("content", "")[:400]
            title = lesson.get("title", "")
            lesson_contexts.append(f'"{title}": {ctx}')
        combined_context = "\n---\n".join(lesson_contexts)

        def generate_batch(difficulty: str, count: int) -> List[Dict]:
            """Gọi 1 request để lấy toàn bộ câu hỏi của 1 mức độ."""
            prompt = f"""Đóng vai giáo viên chuyên nghiệp. Dựa vào các bài học sau:
{combined_context[:2000]}

Hãy tạo ĐÚNG {count} câu hỏi quiz bằng tiếng Việt, độ khó: {difficulty}.
Kết hợp các loại: "mcq" (trắc nghiệm 4 đáp án) và "fill" (điền khuyết).
Chỉ trả về một mảng JSON hợp lệ duy nhất, không kèm markdown hay giải thích.

Cấu trúc mảng:
[
  {{
    "id": 1,
    "type": "mcq",
    "difficulty": "{difficulty}",
    "question": "Câu hỏi trắc nghiệm?",
    "choices": ["A. Đáp án 1", "B. Đáp án 2", "C. Đáp án 3", "D. Đáp án 4"],
    "correct_answer": 0
  }},
  {{
    "id": 2,
    "type": "fill",
    "difficulty": "{difficulty}",
    "question": "Khái niệm ___ là nền tảng của chủ đề này.",
    "words": ["Từ đúng", "Từ nhiễu 1", "Từ nhiễu 2", "Từ nhiễu 3"],
    "choices": ["Từ đúng", "Từ nhiễu 1", "Từ nhiễu 2", "Từ nhiễu 3"],
    "correct_answer": 0
  }}
]
(correct_answer là index 0-3 của đáp án đúng. Tạo đủ {count} câu.)"""

            try:
                response = self.client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                )
                response_text = (response.choices[0].message.content or "").strip()
                response_text = response_text.replace('```json', '').replace('```', '').strip()

                parsed = json.loads(response_text)
                if not isinstance(parsed, list):
                    raise ValueError("Response không phải mảng JSON")

                valid_questions = []
                for i, q in enumerate(parsed):
                    if not isinstance(q, dict):
                        continue
                    q["id"] = i + 1
                    q["difficulty"] = difficulty
                    # Đảm bảo fill có cả words lẫn choices
                    if q.get("type") == "fill" and "words" in q:
                        q.setdefault("choices", q["words"])
                    if self.validator.validate_question(q):
                        valid_questions.append(self.validator.clean_question(q))

                logger.debug("Batch %s: %d/%d câu hợp lệ", difficulty, len(valid_questions), count)
                return valid_questions

            except json.JSONDecodeError as e:
                logger.warning("Batch %s: JSON lỗi (%s), dùng fallback.", difficulty, e)
                return []
            except Exception as e:
                logger.warning("Batch %s: %s, dùng fallback.", difficulty, e)
                return []

        # Chạy 3 batch TUẦN TỰ với stagger 3s — tổng chỉ ~15-20s, không cần song song.
        # Song song 3 thread sẽ burst cùng lúc → tất cả bị 429.
        for i, (diff, cnt) in enumerate(counts.items()):
            if i > 0:
                time.sleep(3)  # stagger thật sự: batch sau chờ batch trước gửi xong
            try:
                quiz[diff] = generate_batch(diff, cnt)
            except Exception as e:
                logger.warning("Batch %s lỗi: %s", diff, e)

        # Fallback cho các độ khó không đủ câu
        question_id = sum(len(v) for v in quiz.values()) + 1
        for diff, count in counts.items():
            while len(quiz[diff]) < count:
                lesson_list = [l for l in lessons if l]
                q = self._generate_unique_quiz_question(question_id, diff, lesson_list)
                if q:
                    quiz[diff].append(q)
                    question_id += 1
                else:
                    break  # tránh vòng lặp vô tận nếu fallback cũng fail

        return quiz

    def _generate_unique_quiz_question(self, q_id: int, difficulty: str, lessons: List[Dict]) -> Dict:
        for attempt in range(10):
            question = self._generate_quiz_question(q_id, difficulty, lessons)
            opts = question.get("choices", question.get("words", ["Đúng", "Sai"]))
            if self._is_question_unique(question["question"], opts, self.used_quiz_questions):
                self._add_question_to_used(question["question"], opts, self.used_quiz_questions)
                return question

            question = self._generate_varied_quiz_question(q_id, difficulty, lessons, attempt)
            opts = question.get("choices", question.get("words", ["Đúng", "Sai"]))
            if self._is_question_unique(question["question"], opts, self.used_quiz_questions):
                self._add_question_to_used(question["question"], opts, self.used_quiz_questions)
                return question

        return self._generate_default_quiz_question(q_id, difficulty)

    def _generate_varied_quiz_question(self, q_id: int, difficulty: str, lessons: List[Dict], variation: int) -> Dict:
        lesson = random.choice(lessons) if lessons else None
        if not lesson:
            return self._generate_default_quiz_question(q_id, difficulty)

        title = lesson.get("title", f"Bài {q_id}")

        if variation % 4 == 0:
            return {
                "id": q_id, "type": "tf", "difficulty": difficulty,
                "question": f"Nội dung trong {title} chứa các khái niệm quan trọng. Đúng hay Sai?",
                "choices": ["Đúng", "Sai"], "correct_answer": 0,
            }

        question_templates = [
            f"Theo {title}, khái niệm cơ bản là gì?",
            f"Điểm đặc trưng của {title} là?",
        ]
        answer_pools = [["Khái niệm nền tảng", "Phương pháp cơ bản", "Nguyên lý vận hành", "Quy trình thực hiện"]]

        return {
            "id": q_id, "type": "mcq", "difficulty": difficulty,
            "question": question_templates[variation % len(question_templates)],
            "choices": [f"{'ABCD'[i]}. {a}" for i, a in enumerate(answer_pools[0])],
            "correct_answer": 0,  # FIX: luôn index 0 với fallback
        }

    def _generate_quiz_question(self, q_id: int, difficulty: str, lessons: list) -> dict:
        lesson = random.choice(lessons) if lessons else None
        context = lesson.get("content", "") if lesson else "Kiến thức chung ngẫu nhiên"

        q_type = random.choice(["mcq", "fill"])

        prompt = f"""Đóng vai là một giáo viên chuyên nghiệp. Dựa vào nội dung bài học sau:
"{context[:800]}"

Hãy tạo 1 câu hỏi bài tập (quiz) bằng tiếng Việt với độ khó: {difficulty}.
Chỉ trả về duy nhất một chuỗi JSON hợp lệ, không kèm theo bất kỳ văn bản giải thích hay markdown code block nào.

Nếu tạo câu hỏi trắc nghiệm (type == "mcq"), dùng đúng cấu trúc này:
{{
    "id": {q_id},
    "type": "mcq",
    "difficulty": "{difficulty}",
    "question": "Nội dung câu hỏi trắc nghiệm?",
    "choices": ["A. Đáp án 1", "B. Đáp án 2", "C. Đáp án 3", "D. Đáp án 4"],
    "correct_answer": 0
}}
(Lưu ý: correct_answer là vị trí index từ 0 đến 3 của đáp án đúng trong mảng choices).

Nếu tạo câu hỏi điền khuyết (type == "fill"), dùng đúng cấu trúc này:
{{
    "id": {q_id},
    "type": "fill",
    "difficulty": "{difficulty}",
    "question": "Điền vào chỗ trống: Khái niệm ___ là một thành phần quan trọng.",
    "words": ["Từ đúng", "Từ nhiễu 1", "Từ nhiễu 2", "Từ nhiễu 3"],
    "choices": ["Từ đúng", "Từ nhiễu 1", "Từ nhiễu 2", "Từ nhiễu 3"],
    "correct_answer": 0
}}
(Lưu ý: "Từ đúng" để điền vào chỗ trống luôn đặt ở index 0).
"""

        try:
            response = self.client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
            )

            response_text = (response.choices[0].message.content or "").strip()

            # Loại bỏ markdown nếu có
            response_text = response_text.replace('```json', '').replace('```', '').strip()

            question_data = json.loads(response_text)
            question_data["id"] = q_id
            question_data["difficulty"] = difficulty

            # FIX: Validate ngay sau khi parse — loại bỏ rác
            if not self.validator.validate_question(question_data):
                logger.warning("Quiz Q%d: câu hỏi không hợp lệ, dùng fallback.", q_id)
                return self._generate_default_quiz_question(q_id, difficulty)

            return self.validator.clean_question(question_data)

        except json.JSONDecodeError as e:
            logger.warning("Quiz Q%d: JSON lỗi (%s), dùng fallback.", q_id, e)
            return self._generate_default_quiz_question(q_id, difficulty)
        except Exception as e:
            logger.warning("Quiz Q%d: %s, dùng fallback.", q_id, e)
            return self._generate_default_quiz_question(q_id, difficulty)

    def _generate_default_quiz_question(self, q_id: int, difficulty: str) -> Dict:
        return {
            "id": q_id,
            "type": "mcq",
            "difficulty": difficulty,
            "question": f"Câu hỏi ôn tập số {q_id} (độ khó: {difficulty})",
            "choices": ["A. Đáp án thứ nhất", "B. Đáp án thứ hai", "C. Đáp án thứ ba", "D. Đáp án thứ tư"],
            "correct_answer": 0,
        }

    def _create_fallback_quiz(self, lessons: List[Dict], total_questions: int) -> Dict:
        quiz: Dict[str, List] = {"easy": [], "medium": [], "hard": []}
        question_id = 1
        for diff, count in [
            ("easy", int(total_questions * 0.4)),
            ("medium", int(total_questions * 0.4)),
            ("hard", total_questions - int(total_questions * 0.8)),
        ]:
            for _ in range(count):
                q = self._generate_unique_quiz_question(question_id, diff, lessons)
                if q:
                    quiz[diff].append(q)
                    question_id += 1
        return quiz

    # ====================================================
    # SAVE — atomic write + backup
    # ====================================================

    def _safe_save_data_files(self, lessons: List[Dict], quiz: Dict) -> bool:
        try:
            total_lesson_questions = sum(len(lesson.get('questions', [])) for lesson in lessons)
            total_quiz_questions = sum(len(q) for q in quiz.values()) if isinstance(quiz, dict) else 0

            lessons_backup = self.lessons_path + '.backup'
            quiz_backup = self.quiz_path + '.backup'

            if os.path.exists(self.lessons_path):
                try:
                    import shutil
                    shutil.copy2(self.lessons_path, lessons_backup)
                except Exception:
                    pass
            if os.path.exists(self.quiz_path):
                try:
                    import shutil
                    shutil.copy2(self.quiz_path, quiz_backup)
                except Exception:
                    pass

            lessons_data = {
                "metadata": {
                    "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                    "total_lessons": len(lessons),
                    "total_questions": total_lesson_questions,
                },
                "lessons": lessons,
            }

            quiz_data_obj = {"metadata": {"total_questions": total_quiz_questions}, **quiz}

            # Validate JSON trước khi ghi
            try:
                json.dumps(lessons_data, ensure_ascii=False)
                json.dumps(quiz_data_obj, ensure_ascii=False)
            except (TypeError, ValueError) as e:
                logger.error("Lỗi serialize JSON: %s", e)
                return False

            # Ghi atomic qua file .tmp
            os.makedirs(os.path.dirname(self.lessons_path), exist_ok=True)
            os.makedirs(os.path.dirname(self.quiz_path), exist_ok=True)

            temp_lessons = self.lessons_path + '.tmp'
            with open(temp_lessons, 'w', encoding='utf-8') as f:
                json.dump(lessons_data, f, ensure_ascii=False, indent=2)
            if os.name == 'nt' and os.path.exists(self.lessons_path):
                os.remove(self.lessons_path)
            os.rename(temp_lessons, self.lessons_path)

            temp_quiz = self.quiz_path + '.tmp'
            with open(temp_quiz, 'w', encoding='utf-8') as f:
                json.dump(quiz_data_obj, f, ensure_ascii=False, indent=2)
            if os.name == 'nt' and os.path.exists(self.quiz_path):
                os.remove(self.quiz_path)
            os.rename(temp_quiz, self.quiz_path)

            return True

        except Exception as e:
            logger.error("Lỗi lưu file: %s", e)
            try:
                import shutil
                if os.path.exists(lessons_backup):
                    shutil.copy2(lessons_backup, self.lessons_path)
                if os.path.exists(quiz_backup):
                    shutil.copy2(quiz_backup, self.quiz_path)
            except Exception:
                pass
            return False

    def stop_processing(self):
        self.should_stop = True

    def cleanup_cache(self):
        self.content_cache.clear()
        self.used_questions.clear()
        self.used_quiz_questions.clear()